#!/usr/bin/python3
# _*_ coding:utf-8 _*_
#
# Copyright (C) 2023 - 2023 pgl, Inc. All Rights Reserved 
#
# @Time    : 2023/2/gyx 9:21
# @Author  : pgl
# @File    : 02_读取word文件.py.py
# @IDE     : PyCharm


from docx import Document

doc = Document('./res/用函数还是用复杂的表达式.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
# print(doc.paragraphs[1].runs[0].text)

content = []
for para in doc.paragraphs:
    content.append(para.text)
print(''.join(content))
